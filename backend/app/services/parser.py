"""
CV text extraction service.

Extracts raw text from PDF, DOCX, and DOC files.
"""
import os
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

    return text.strip()


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Also extract text from tables (CVs often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    return text.strip()


def extract_text_from_doc(file_path):
    """Extract text from a legacy DOC file by converting to DOCX first.
    
    Requires LibreOffice installed on the system.
    Falls back to basic binary text extraction if LibreOffice is not available.
    """
    import subprocess

    # Try converting with LibreOffice
    output_dir = os.path.dirname(file_path)
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", "--outdir", output_dir, file_path],
            capture_output=True,
            timeout=30,
            check=True,
        )
        # The converted file will have the same name but .docx extension
        docx_path = file_path.rsplit(".", 1)[0] + ".docx"
        if os.path.exists(docx_path):
            text = extract_text_from_docx(docx_path)
            os.remove(docx_path)  # Clean up the converted file
            return text
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: basic binary text extraction for .doc files
    try:
        with open(file_path, "rb") as f:
            raw = f.read()
        # Extract printable ASCII text chunks from the binary
        text = ""
        current_chunk = ""
        for byte in raw:
            if 32 <= byte <= 126 or byte in (10, 13):
                current_chunk += chr(byte)
            else:
                if len(current_chunk) > 20:  # Only keep meaningful chunks
                    text += current_chunk + "\n"
                current_chunk = ""
        if len(current_chunk) > 20:
            text += current_chunk
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from DOC: {str(e)}")


def extract_text(file_path, file_format):
    """Extract text from a CV file based on its format.
    
    Args:
        file_path: Path to the uploaded file
        file_format: File extension (pdf, docx, doc)
    
    Returns:
        Extracted text as a string
    """
    if file_format == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_format == "docx":
        return extract_text_from_docx(file_path)
    elif file_format == "doc":
        return extract_text_from_doc(file_path)
    else:
        raise Exception(f"Unsupported file format: {file_format}")
